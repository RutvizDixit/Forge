from __future__ import annotations

import json
import re
import uuid
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import urlparse

import pandas as pd
import requests
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

import config


@dataclass
class IngestedSource:
    source_id: str
    source_name: str
    source_type: str
    status: str
    source_path: Optional[str] = None
    source_url: Optional[str] = None
    rows: List[Dict[str, Any]] = field(default_factory=list)
    columns: List[str] = field(default_factory=list)
    text: str = ""
    children: List[Dict[str, Any]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    error: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


class SourceIngestor:
    """
    Universal source ingestion layer for FORGE.

    Supported sources:

    CSV
    XLSX
    XLS
    PDF
    DOCX
    TXT
    JSON
    XML
    ZIP
    Website URLs

    The ingestor extracts source material. It does not decide whether
    extracted product information is factually correct.
    """

    REQUEST_TIMEOUT = 25

    USER_AGENT = (
        "FORGE/1.0 "
        "(Industrial Product Intelligence; source ingestion)"
    )

    def ingest_file(
        self,
        file_path: str | Path,
    ) -> IngestedSource:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(
                f"Source file does not exist: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Source path is not a file: {path}"
            )

        extension = path.suffix.lower().lstrip(".")

        if extension not in config.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: .{extension or 'unknown'}"
            )

        source_id = self._source_id()

        base_metadata = {
            "filename": path.name,
            "extension": extension,
            "size_bytes": path.stat().st_size,
        }

        try:
            if extension in {"csv", "xlsx", "xls"}:
                result = self._ingest_table(
                    path,
                    source_id,
                )

            elif extension == "pdf":
                result = self._ingest_pdf(
                    path,
                    source_id,
                )

            elif extension == "docx":
                result = self._ingest_docx(
                    path,
                    source_id,
                )

            elif extension == "txt":
                result = self._ingest_text(
                    path,
                    source_id,
                )

            elif extension == "json":
                result = self._ingest_json(
                    path,
                    source_id,
                )

            elif extension == "xml":
                result = self._ingest_xml(
                    path,
                    source_id,
                )

            elif extension == "zip":
                result = self._ingest_zip(
                    path,
                    source_id,
                )

            else:
                raise ValueError(
                    f"No ingestion handler exists for .{extension}"
                )

            result.metadata.update(base_metadata)

            return result

        except Exception as exc:
            return IngestedSource(
                source_id=source_id,
                source_name=path.name,
                source_type=self._classify_extension(extension),
                status="failed",
                source_path=str(path),
                metadata=base_metadata,
                error=str(exc),
            )

    def ingest_files(
        self,
        file_paths: List[str | Path],
    ) -> List[IngestedSource]:
        results = []

        for file_path in file_paths:
            results.append(
                self.ingest_file(file_path)
            )

        return results

    def ingest_url(
        self,
        url: str,
    ) -> IngestedSource:
        source_id = self._source_id()

        normalized_url = self._validate_url(url)

        try:
            response = requests.get(
                normalized_url,
                timeout=self.REQUEST_TIMEOUT,
                headers={
                    "User-Agent": self.USER_AGENT,
                },
                allow_redirects=True,
            )

            response.raise_for_status()

            content_type = (
                response.headers.get(
                    "Content-Type",
                    "",
                )
                .lower()
            )

            final_url = response.url

            if (
                "text/html" not in content_type
                and not content_type.startswith(
                    "application/xhtml"
                )
            ):
                raise ValueError(
                    "The supplied URL did not return an HTML page."
                )

            text, page_metadata = self._extract_web_page(
                response.text,
                final_url,
            )

            return IngestedSource(
                source_id=source_id,
                source_name=self._web_source_name(
                    final_url
                ),
                source_type="Website",
                status="processed",
                source_url=final_url,
                text=text,
                metadata={
                    "content_type": content_type,
                    "status_code": response.status_code,
                    **page_metadata,
                },
            )

        except Exception as exc:
            return IngestedSource(
                source_id=source_id,
                source_name=self._web_source_name(
                    normalized_url
                ),
                source_type="Website",
                status="failed",
                source_url=normalized_url,
                error=str(exc),
            )

    def _ingest_table(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        extension = path.suffix.lower()

        if extension == ".csv":
            dataframe = self._read_csv(path)
        else:
            dataframe = self._read_excel(
                path
            )

        dataframe = dataframe.fillna("")

        rows = dataframe.to_dict(
            orient="records"
        )

        rows = [
            self._clean_record(row)
            for row in rows
        ]

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Catalogue",
            status="processed",
            source_path=str(path),
            rows=rows,
            columns=[
                str(column)
                for column in dataframe.columns
            ],
            metadata={
                "record_count": len(rows),
                "field_count": len(
                    dataframe.columns
                ),
            },
        )

    def _ingest_pdf(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        reader = PdfReader(str(path))

        page_text = []
        page_metadata = []

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            text = page.extract_text() or ""

            text = self._clean_text(text)

            page_text.append(text)

            page_metadata.append(
                {
                    "page": page_number,
                    "characters": len(text),
                }
            )

        full_text = "\n\n".join(
            f"[Page {item['page']}]\n"
            f"{page_text[item['page'] - 1]}"
            for item in page_metadata
            if page_text[item["page"] - 1]
        )

        warnings = []

        if not full_text.strip():
            warnings.append(
                "No extractable text was found in the PDF."
            )

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Technical document",
            status="processed",
            source_path=str(path),
            text=full_text,
            metadata={
                "page_count": len(reader.pages),
                "pages": page_metadata,
            },
            warnings=warnings,
        )

    def _ingest_docx(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        document = Document(str(path))

        paragraphs = []

        for paragraph in document.paragraphs:
            text = self._clean_text(
                paragraph.text
            )

            if text:
                paragraphs.append(text)

        table_data = []

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            rows = []

            for row in table.rows:
                rows.append(
                    [
                        self._clean_text(
                            cell.text
                        )
                        for cell in row.cells
                    ]
                )

            table_data.append(
                {
                    "table": table_index,
                    "rows": rows,
                }
            )

        text_parts = paragraphs[:]

        for table in table_data:
            text_parts.append(
                self._table_to_text(table)
            )

        full_text = "\n\n".join(
            text_parts
        )

        warnings = []

        if not full_text.strip():
            warnings.append(
                "No readable text or table content was found in the DOCX."
            )

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Product document",
            status="processed",
            source_path=str(path),
            text=full_text,
            metadata={
                "paragraph_count": len(
                    paragraphs
                ),
                "table_count": len(
                    table_data
                ),
                "tables": table_data,
            },
            warnings=warnings,
        )

    def _ingest_text(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        text = path.read_text(
            encoding="utf-8",
            errors="replace",
        )

        text = self._clean_text(text)

        warnings = []

        if not text:
            warnings.append(
                "The text source is empty."
            )

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Text source",
            status="processed",
            source_path=str(path),
            text=text,
            metadata={
                "characters": len(text),
            },
            warnings=warnings,
        )

    def _ingest_json(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        raw_text = path.read_text(
            encoding="utf-8",
            errors="replace",
        )

        payload = json.loads(raw_text)

        rows = self._json_to_rows(payload)

        text = json.dumps(
            payload,
            indent=2,
            ensure_ascii=False,
        )

        columns = []

        if rows:
            columns = list(
                rows[0].keys()
            )

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Structured data",
            status="processed",
            source_path=str(path),
            rows=rows,
            columns=columns,
            text=text,
            metadata={
                "record_count": len(rows),
                "json_root_type": type(
                    payload
                ).__name__,
            },
        )

    def _ingest_xml(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        raw_text = path.read_text(
            encoding="utf-8",
            errors="replace",
        )

        soup = BeautifulSoup(
            raw_text,
            "xml",
        )

        text = self._clean_text(
            soup.get_text(
                " ",
                strip=True,
            )
        )

        rows = []

        for element in soup.find_all(
            recursive=True
        ):
            children = [
                child
                for child in element.find_all(
                    recursive=False
                )
                if child.name
            ]

            if not children:
                continue

            record = {}

            for child in children:
                value = self._clean_text(
                    child.get_text(
                        " ",
                        strip=True,
                    )
                )

                if value:
                    record[
                        child.name
                    ] = value

            if record:
                rows.append(record)

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Structured data",
            status="processed",
            source_path=str(path),
            rows=rows,
            columns=(
                list(rows[0].keys())
                if rows
                else []
            ),
            text=text,
            metadata={
                "record_count": len(rows),
            },
        )

    def _ingest_zip(
        self,
        path: Path,
        source_id: str,
    ) -> IngestedSource:
        children = []
        warnings = []

        extraction_root = (
            config.UPLOAD_DIR
            / f"archive_{source_id}"
        )

        extraction_root.mkdir(
            parents=True,
            exist_ok=True,
        )

        with zipfile.ZipFile(
            path,
            "r",
        ) as archive:
            members = archive.infolist()

            safe_members = []

            for member in members:
                if member.is_dir():
                    continue

                target = (
                    extraction_root
                    / member.filename
                )

                if not self._safe_archive_target(
                    extraction_root,
                    target,
                ):
                    warnings.append(
                        f"Skipped unsafe archive path: "
                        f"{member.filename}"
                    )
                    continue

                extension = (
                    Path(
                        member.filename
                    )
                    .suffix
                    .lower()
                    .lstrip(".")
                )

                if extension not in config.ALLOWED_EXTENSIONS:
                    warnings.append(
                        f"Skipped unsupported archive file: "
                        f"{member.filename}"
                    )
                    continue

                safe_members.append(
                    member
                )

            archive.extractall(
                extraction_root,
                members=safe_members,
            )

        for member in safe_members:
            child_path = (
                extraction_root
                / member.filename
            )

            try:
                child = self.ingest_file(
                    child_path
                )

                children.append(
                    child.to_dict()
                )

            except Exception as exc:
                warnings.append(
                    f"Could not process "
                    f"{member.filename}: {exc}"
                )

        text_parts = []

        combined_rows = []
        combined_columns = []

        for child in children:
            child_text = child.get(
                "text",
                "",
            )

            if child_text:
                text_parts.append(
                    f"[{child['source_name']}]\n"
                    f"{child_text}"
                )

            rows = child.get(
                "rows",
                [],
            )

            if rows:
                combined_rows.extend(
                    rows
                )

            for column in child.get(
                "columns",
                [],
            ):
                if column not in combined_columns:
                    combined_columns.append(
                        column
                    )

        return IngestedSource(
            source_id=source_id,
            source_name=path.name,
            source_type="Source bundle",
            status="processed",
            source_path=str(path),
            rows=combined_rows,
            columns=combined_columns,
            text="\n\n".join(
                text_parts
            ),
            children=children,
            metadata={
                "archive_file_count": len(
                    safe_members
                ),
                "processed_file_count": len(
                    children
                ),
                "extraction_directory": str(
                    extraction_root
                ),
            },
            warnings=warnings,
        )

    def _read_csv(
        self,
        path: Path,
    ) -> pd.DataFrame:
        encodings = (
            "utf-8-sig",
            "utf-8",
            "cp1252",
            "latin1",
        )

        last_error = None

        for encoding in encodings:
            try:
                return pd.read_csv(
                    path,
                    dtype=str,
                    encoding=encoding,
                    keep_default_na=False,
                )
            except UnicodeDecodeError as exc:
                last_error = exc

        raise ValueError(
            f"Could not decode CSV file: {last_error}"
        )

    def _read_excel(
        self,
        path: Path,
    ) -> pd.DataFrame:
        return pd.read_excel(
            path,
            dtype=str,
        ).fillna("")

    def _extract_web_page(
        self,
        html: str,
        url: str,
    ):
        soup = BeautifulSoup(
            html,
            "html.parser",
        )

        title = ""

        if soup.title:
            title = self._clean_text(
                soup.title.get_text(
                    " ",
                    strip=True,
                )
            )

        meta_description = ""

        meta = soup.find(
            "meta",
            attrs={
                "name": re.compile(
                    "^description$",
                    re.IGNORECASE,
                )
            },
        )

        if meta:
            meta_description = self._clean_text(
                meta.get(
                    "content",
                    "",
                )
            )

        for element in soup(
            [
                "script",
                "style",
                "noscript",
                "svg",
                "template",
            ]
        ):
            element.decompose()

        text = self._clean_text(
            soup.get_text(
                " ",
                strip=True,
            )
        )

        return text, {
            "title": title,
            "description": meta_description,
            "domain": urlparse(url).netloc,
        }

    def _validate_url(
        self,
        url: str,
    ) -> str:
        url = str(url).strip()

        if not url:
            raise ValueError(
                "Website URL cannot be empty."
            )

        parsed = urlparse(url)

        if parsed.scheme not in {
            "http",
            "https",
        }:
            raise ValueError(
                "Website URL must begin with http:// or https://."
            )

        if not parsed.netloc:
            raise ValueError(
                "Website URL does not contain a valid host."
            )

        if config.ALLOWED_WEB_HOSTS:
            host = (
                parsed.hostname
                or ""
            ).lower()

            if not any(
                host == allowed
                or host.endswith(
                    f".{allowed}"
                )
                for allowed in config.ALLOWED_WEB_HOSTS
            ):
                raise ValueError(
                    "This website host is not allowed by FORGE configuration."
                )

        return url

    @staticmethod
    def _classify_extension(
        extension: str,
    ) -> str:
        return config.SUPPORTED_SOURCE_TYPES.get(
            extension,
            "Unknown",
        )

    @staticmethod
    def _source_id() -> str:
        return (
            f"src-{uuid.uuid4().hex[:12]}"
        )

    @staticmethod
    def _clean_text(
        text: Any,
    ) -> str:
        if text is None:
            return ""

        return re.sub(
            r"\s+",
            " ",
            str(text),
        ).strip()

    @staticmethod
    def _clean_record(
        record: Dict[str, Any],
    ) -> Dict[str, Any]:
        cleaned = {}

        for key, value in record.items():
            if value is None:
                value = ""

            if isinstance(
                value,
                float,
            ) and pd.isna(value):
                value = ""

            cleaned[str(key)] = str(
                value
            ).strip()

        return cleaned

    @staticmethod
    def _json_to_rows(
        payload: Any,
    ) -> List[Dict[str, Any]]:
        if isinstance(
            payload,
            list,
        ):
            rows = []

            for item in payload:
                if isinstance(
                    item,
                    dict,
                ):
                    rows.append(
                        SourceIngestor._clean_record(
                            item
                        )
                    )
                else:
                    rows.append(
                        {
                            "value": item
                        }
                    )

            return rows

        if isinstance(
            payload,
            dict,
        ):
            for key in (
                "data",
                "items",
                "products",
                "records",
                "results",
            ):
                value = payload.get(
                    key
                )

                if isinstance(
                    value,
                    list,
                ):
                    return SourceIngestor._json_to_rows(
                        value
                    )

            return [
                SourceIngestor._clean_record(
                    payload
                )
            ]

        return [
            {
                "value": payload
            }
        ]

    @staticmethod
    def _table_to_text(
        table: Dict[str, Any],
    ) -> str:
        lines = [
            f"Table {table['table']}:"
        ]

        for row in table["rows"]:
            lines.append(
                " | ".join(row)
            )

        return "\n".join(lines)

    @staticmethod
    def _safe_archive_target(
        root: Path,
        target: Path,
    ) -> bool:
        try:
            root_resolved = root.resolve()
            target_resolved = target.resolve()

            target_resolved.relative_to(
                root_resolved
            )

            return True

        except ValueError:
            return False

    @staticmethod
    def _web_source_name(
        url: str,
    ) -> str:
        parsed = urlparse(url)

        return (
            parsed.netloc
            or url
        )